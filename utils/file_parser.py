"""
文件解析工具
支持多种文件格式转换为文本
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
import mimetypes


class FileParser:
    """文件解析器 - 支持多种格式"""
    
    # 支持的文件类型
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.markdown': 'text/markdown',
        '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.json': 'application/json',
        '.csv': 'text/csv'
    }
    
    def __init__(self):
        """初始化解析器"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查可用的解析库"""
        self.has_pypdf2 = False
        self.has_pdfplumber = False
        self.has_python_docx = False
        self.has_beautifulsoup = False
        
        try:
            import PyPDF2
            self.has_pypdf2 = True
        except ImportError:
            pass
        
        try:
            import pdfplumber
            self.has_pdfplumber = True
        except ImportError:
            pass
        
        try:
            import docx
            self.has_python_docx = True
        except ImportError:
            pass
        
        try:
            from bs4 import BeautifulSoup
            self.has_beautifulsoup = True
        except ImportError:
            pass
    
    def is_supported(self, filename: str) -> bool:
        """检查文件是否支持"""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        return self.SUPPORTED_EXTENSIONS.get(ext)
    
    def parse_file(
        self,
        file_path: str,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        解析文件并提取文本
        
        Args:
            file_path: 文件路径
            filename: 原始文件名（可选）
            
        Returns:
            包含文本和元数据的字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 确定文件类型
        if filename:
            ext = Path(filename).suffix.lower()
        else:
            ext = Path(file_path).suffix.lower()
        
        if not self.is_supported(filename or file_path):
            raise ValueError(f"不支持的文件格式: {ext}")
        
        # 根据文件类型选择解析方法
        parsers = {
            '.txt': self._parse_txt,
            '.md': self._parse_txt,
            '.markdown': self._parse_txt,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.json': self._parse_json,
            '.csv': self._parse_csv
        }
        
        parser_func = parsers.get(ext)
        if not parser_func:
            raise ValueError(f"暂不支持的文件格式: {ext}")
        
        try:
            result = parser_func(file_path)
            result['file_type'] = ext
            result['file_name'] = filename or os.path.basename(file_path)
            result['file_size'] = os.path.getsize(file_path)
            return result
        except Exception as e:
            raise RuntimeError(f"解析文件失败: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """解析纯文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                return {
                    'content': content,
                    'encoding': encoding,
                    'char_count': len(content),
                    'line_count': content.count('\n') + 1
                }
            except (UnicodeDecodeError, Exception):
                continue
        
        raise ValueError("无法识别文件编码")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文件"""
        # 优先使用 pdfplumber（更准确）
        if self.has_pdfplumber:
            return self._parse_pdf_with_pdfplumber(file_path)
        elif self.has_pypdf2:
            return self._parse_pdf_with_pypdf2(file_path)
        else:
            raise ImportError(
                "需要安装 PDF 解析库: pip install pdfplumber 或 pip install PyPDF2"
            )
    
    def _parse_pdf_with_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """使用 pdfplumber 解析 PDF"""
        import pdfplumber
        
        text_parts = []
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata['page_count'] = len(pdf.pages)
            metadata['metadata'] = pdf.metadata
            
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        content = '\n\n'.join(text_parts)
        
        return {
            'content': content,
            'char_count': len(content),
            **metadata
        }
    
    def _parse_pdf_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """使用 PyPDF2 解析 PDF"""
        import PyPDF2
        
        text_parts = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            page_count = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        content = '\n\n'.join(text_parts)
        
        return {
            'content': content,
            'char_count': len(content),
            'page_count': page_count
        }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档 (.docx)"""
        if not self.has_python_docx:
            raise ImportError("需要安装: pip install python-docx")
        
        import docx
        
        doc = docx.Document(file_path)
        
        # 提取段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 提取表格文本
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                tables_text.append(row_text)
        
        content = '\n\n'.join(paragraphs)
        if tables_text:
            content += '\n\n' + '\n'.join(tables_text)
        
        return {
            'content': content,
            'char_count': len(content),
            'paragraph_count': len(paragraphs),
            'table_count': len(doc.tables)
        }
    
    def _parse_html(self, file_path: str) -> Dict[str, Any]:
        """解析 HTML 文件"""
        if not self.has_beautifulsoup:
            raise ImportError("需要安装: pip install beautifulsoup4")
        
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除 script 和 style 标签
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 提取文本
        text = soup.get_text()
        
        # 清理多余的空白
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return {
            'content': text,
            'char_count': len(text),
            'title': soup.title.string if soup.title else None
        }
    
    def _parse_json(self, file_path: str) -> Dict[str, Any]:
        """解析 JSON 文件"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 将 JSON 转为格式化的文本
        content = json.dumps(data, ensure_ascii=False, indent=2)
        
        return {
            'content': content,
            'char_count': len(content),
            'json_data': data
        }
    
    def _parse_csv(self, file_path: str) -> Dict[str, Any]:
        """解析 CSV 文件"""
        import csv
        
        rows = []
        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.reader(f)
            for row in csv_reader:
                rows.append(' | '.join(row))
        
        content = '\n'.join(rows)
        
        return {
            'content': content,
            'char_count': len(content),
            'row_count': len(rows)
        }
    
    def get_supported_formats(self) -> Dict[str, str]:
        """获取支持的格式列表"""
        formats = {}
        
        formats['.txt'] = '纯文本'
        formats['.md'] = 'Markdown'
        
        if self.has_pdfplumber or self.has_pypdf2:
            formats['.pdf'] = 'PDF 文档'
        
        if self.has_python_docx:
            formats['.docx'] = 'Word 文档'
        
        if self.has_beautifulsoup:
            formats['.html'] = 'HTML 文档'
        
        formats['.json'] = 'JSON 文件'
        formats['.csv'] = 'CSV 文件'
        
        return formats


# 全局解析器实例
file_parser = FileParser()


def parse_uploaded_file(file_path: str, filename: str) -> Dict[str, Any]:
    """
    解析上传的文件
    
    Args:
        file_path: 文件保存路径
        filename: 原始文件名
        
    Returns:
        包含解析结果的字典
    """
    return file_parser.parse_file(file_path, filename)


def get_supported_file_types() -> Dict[str, str]:
    """获取支持的文件类型"""
    return file_parser.get_supported_formats()


if __name__ == "__main__":
    # 测试代码
    parser = FileParser()
    print("支持的文件格式:")
    for ext, name in parser.get_supported_formats().items():
        print(f"  {ext}: {name}")

